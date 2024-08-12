import fitz  # PyMuPDF
import nltk
import streamlit as st
from docx import Document
import openai
import os
from nltk.tokenize import sent_tokenize
import pandas as pd

openai.api_key = st.secrets["OpenAI_API_key"]

# Initialize session state for variables that need to persist
if "FirstList" not in st.session_state:
    st.session_state.FirstList = []

if "edit_index" not in st.session_state:
    st.session_state.edit_index = None  # No entry is being edited initially

def extract_and_split_text(file_path: str, chunk_size: int = 500) -> list:
    def extract_text_from_docx(docx_path: str) -> str:
        doc = Document(docx_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text

    if file_path.endswith('.pdf'):
        # Open the PDF file
        document = fitz.open(file_path)
        text = ""
        # Iterate through the pages
        for page_num in range(document.page_count):
            page = document.load_page(page_num)
            text += page.get_text()
    elif file_path.endswith('.docx'):
        text = extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file format. Please provide a PDF or DOCX file.")

    # Split text into sentences
    sentences = sent_tokenize(text)
    chunks = []
    current_chunk = ""

    for sentence in sentences:
        if len(current_chunk) + len(sentence) <= chunk_size:
            current_chunk += sentence + " "
        else:
            chunks.append(current_chunk.strip())
            current_chunk = sentence + " "

    if current_chunk:
        chunks.append(current_chunk.strip())

    return chunks


def ListofTerms(lang, Text):
    prompt = f'''<role> 
            You are a professional translator. You will help the user extract all the technical words from a text and translate them to {lang}.
            <\\role>

            <rules>
            -Only return a list of the most common technical terms in the text.
            -DO NOT try to translate the whole document
            -If a term is present multiple times, only return one instance of the term. DO NOT REPEAT words.
            -DO NOT translate non technical words or common language words.
            -If the text is of a legal nature, check the web for Professional Translations of the legal terms
            -DO NOT number the final list
            -If a word has more than one translation, return only the most technical or professional one. If you have trouble finding it, check the web.
            -The input language is not always english. Detect the language and understand the document well before tarting the extraction and translation.
            <\\rules>

            <example>
            Language to be translated into: french

            <Text>
            I am writing to apply for the position of machine operator in the Pastilab factory, as advertised by the company's LinkedIn page. I have included my resume for your consideration.
            I am a student at the Lebanese university, where I am majoring in electrical engineering at the Roumieh campus. In addition to my education, I have participated in a multitude of mechatronics workshops that introduced me to the world of machines and operation. And I believe the closeness of the campus to the factory grounds make me a suitable fit for the company.
            Thank you for your time and your consideration. I hope to have compelled you to accept me as an employee at the leading company in the plastics industry in Lebanon,
            <\\Text>

            <expected output>
            Machine operator: opérateur de machine
            Resume: CV
            Electrical engineering: Genie Electrique
            Mechatronics: mécatronique
            Factory grounds: terrain d'usine
            <\\expected output>

            <explanation>
            This is an example of a cover page sent for a job application. The output was only the technical terms that could lose their meaning in translation.
            <\\explanation> 

            <\\example>

            '''

    response = openai.ChatCompletion.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": prompt},
            {"role": "user", "content": Text},
        ],
        temperature=0,
    )
    output_text = response.choices[0].message['content']
    
    # Splitting the text into a list
    translation_list = [entry.split(':') for entry in output_text.split('\n')]
    return translation_list

st.title('Document Translator')

# File uploader
uploaded_file = st.file_uploader("Choose a DOCX or PDF file", type=["docx", "pdf"])

if uploaded_file is not None:
    # Read the file content correctly
    doc = uploaded_file.name
        
    # Chunking text
    chunked_text = extract_and_split_text(doc, 600)
    full_text = '\n'.join(chunked_text)

    # Language input
    language = st.text_input("Desired output language")

    if st.button('Generate Technical Terms List'):
        # Generate the list of technical terms
       for i in range (0,4): 
        try:
            st.session_state.FirstList = ListofTerms(language, full_text)
            break
        except:
            continue
    
    # Display the DataFrame as a table
    if st.session_state.FirstList:
        df = pd.DataFrame(st.session_state.FirstList, columns=['original term', 'proposed translation'])
        st.dataframe(df)
    
    # Edit table logic
    st.write("Use the following inputs to edit the table:")

    # Add entry
    OGTerm = st.text_input('Original term', key="add_term")
    PTrans = st.text_input('Proposed Translation', key="add_translation")
    if st.button('Add entry'):
        if OGTerm and PTrans:
            st.session_state.FirstList.append([OGTerm, PTrans])
    
    # Remove entry
    entry_to_remove = st.number_input('Select which number entry you want to delete', min_value=0, max_value=len(st.session_state.FirstList)-1, key="remove_entry")
    if st.button('Remove entry'):
        if 0 <= entry_to_remove < len(st.session_state.FirstList):
            st.session_state.FirstList.pop(entry_to_remove)
    
    # Edit entry
    entry_to_edit = st.number_input('Select which number entry you want to edit', min_value=0, max_value=len(st.session_state.FirstList)-1, key="edit_entry")
    if st.button('Edit entry'):
        if 0 <= entry_to_edit < len(st.session_state.FirstList):
            st.session_state.edit_index = entry_to_edit  # Store the index of the entry being edited

    # If an entry is selected for editing, provide input for the new translation
    if st.session_state.edit_index is not None:
        NewTranslation = st.text_input('New proposed translation:', key="edit_translation")
        if st.button('Confirm Edit'):
            if NewTranslation:
                st.session_state.FirstList[st.session_state.edit_index][1] = NewTranslation
                st.session_state.edit_index = None  # Reset after editing

    try:
        ListToString = '\n'.join([': '.join(entry) for entry in st.session_state.FirstList if isinstance(entry, list) and len(entry) == 2])
    except Exception as e:
        st.error(f"Error occurred while processing the list: {e}")
        st.stop()

    # Proceed with translation
    if st.button('Proceed with translation'):
        for i in range(0,4):
            try:
                firstTranslation = openai.ChatCompletion.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": f'''<role>
                            You are a professional translator. You will be given a chunk of text that you need to translate into {language}. Follow professional translating principles.
                            <\\role>

                            <rules>
                            -The translation should be word for word.
                            -DO NOT alter the meaning of the text in any way shape or form.
                            -DO NOT translate links under any circumstance.
                            -Should you encounter a technical term, refer to this list: {ListToString}
                            -Only use this list to translate the terms that are mentioned in said list.
                            -The word count of the text should remain roughly the same.
                            -DO NOT expand or contract on ideas, sentences or phrases.
                            -Read the text well and understand its context before starting to translate.
                            <\\rules>
                            '''},
                        {"role": "user", "content": chunked_text[0]},
                    ],
                    temperature=0,
                )
                break
            except:
                continue
        progress_bar = st.progress(0)
        Translation = Document()
        # Add the GPT text to the new doc
        Translation.add_paragraph(firstTranslation.choices[0].message['content'] + '\n\n')
        # Save the new document
        Translation.save('translation.docx')
        for i in range(1, len(chunked_text)):
            ExistingTextList = [para.text for para in Translation.paragraphs]
            ExistingText = "\n".join(ExistingTextList)
            for j in range(0,4):
                try:
                    Loop = openai.ChatCompletion.create(
                        model="gpt-4o-mini",
                        messages=[
                            {"role": "system", "content": f'''<role>
                                You are a professional translator. You will be given a chunk of text that you need to translate into {language}. Follow professional translating principles.
                                <\\role>

                                <rules>
                                -The translation should be word for word.
                                -DO NOT alter the meaning of the text in any way shape or form.
                                -DO NOT translate links under any circumstance.
                                -Should you encounter a technical term, refer to this list: {ListToString}
                                -Only use this list to translate the terms that are mentioned in said list.
                                -The word count of the text should remain roughly the same.
                                -DO NOT expand or contract on ideas, sentences or phrases.
                                -Read the text well and understand its context before starting to translate.
                                <\\rules>

                                Here are the already translated parts of the text. Read them well, understand the context and follow in the steps of the translation for the new chunk of text.

                                <already translated text>
                                {ExistingText}
                                <\\already translated text>

                                '''},
                            {"role": "user", "content": chunked_text[i]},
                        ],
                        temperature=0,
                    )
                    break
                except:
                    continue
            
            NewChunk = Loop.choices[0].message['content']
            Translation.add_paragraph(NewChunk + '\n\n')
            Translation.save('translation.docx')
            # Update progress bar
            progress = (i + 1) / len(chunked_text)
            progress_bar.progress(progress)
        st.write("Translation completed. Download the file below.")
        with open("translation.docx", "rb") as file:
            st.download_button(
                label="Download Translation",
                data=file,
                file_name="translation.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
